from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "result-report-draft.md"
OUTPUT = ROOT / "docs" / "LGTM_Observability_Stack_결과보고서.docx"

CONTENT_WIDTH_IN = 6.5
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 31, 43)
MUTED = RGBColor(94, 104, 117)
TABLE_HEADER_FILL = "F2F4F7"
CODE_FILL = "F6F8FA"
FONT = "Malgun Gothic"
MONO_FONT = "Consolas"


def set_run_font(run, name=FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def set_para_border_bottom(paragraph, color="D7DBE2", size="8", space="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def setup_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, side, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
        ("Heading 4", 11, DARK_BLUE, 6, 3),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "LGTM Observability Stack 결과보고서"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.runs[0], size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Page ")
    set_run_font(r, size=9, color=MUTED)
    add_page_number(footer)


def add_title_page(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("RESULT REPORT")
    set_run_font(run, size=11, color=MUTED, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(title.replace(" 초안", ""))
    set_run_font(run, size=24, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("Loki · Grafana · Tempo · Mimir 기반 2-VM 관측성 스택 구축")
    set_run_font(run, size=12, color=MUTED)

    meta = [
        ("프로젝트", "LGTM Observability Stack 구축"),
        ("구성", "Monitoring VM + App VM / Docker Compose + K3S"),
        ("주요 개선", "Promtail 기반 초기 구조에서 Alloy 기반 통합 수집 구조로 고도화"),
        ("작성일", "2026년 7월"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Table Grid"
    set_table_width(table)
    set_cell_margins(table)
    for row, (label, value) in zip(table.rows, meta):
        prevent_row_split(row)
        row.cells[0].width = Inches(1.35)
        row.cells[1].width = Inches(5.15)
        set_cell_shading(row.cells[0], TABLE_HEADER_FILL)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
        set_run_font(row.cells[0].paragraphs[0].add_run(label), bold=True)
        set_run_font(row.cells[1].paragraphs[0].add_run(value))

    rule = doc.add_paragraph()
    set_para_border_bottom(rule)
    doc.add_page_break()


def add_toc(doc, headings):
    p = doc.add_paragraph()
    p.style = "Heading 1"
    p.add_run("목차")
    for level, text in headings:
        if level > 2:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25 if level == 2 else 0)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        set_run_font(run, size=10.5 if level == 2 else 11, color=INK, bold=(level == 1))
    doc.add_page_break()


def split_table_row(line):
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def is_table_separator(line):
    return bool(re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", line))


def add_inline_runs(paragraph, text, bold_default=False):
    # Basic inline formatter for `code`, **bold**, and [text](url). It keeps
    # punctuation plain and avoids fragile full Markdown parsing.
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, bold=bold_default)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, MONO_FONT, size=10, color=RGBColor(80, 80, 80))
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, bold=True)
        else:
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            run = paragraph.add_run(m.group(1))
            set_run_font(run, color=BLUE)
            run.underline = True
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, bold=bold_default)


def add_code_block(doc, lines):
    for idx, line in enumerate(lines):
        p = doc.add_paragraph()
        p.style = "No Spacing"
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.right_indent = Inches(0.1)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.keep_together = True
        p.paragraph_format.keep_with_next = idx < len(lines) - 1
        run = p.add_run(line if line else " ")
        set_run_font(run, MONO_FONT, size=8.5, color=RGBColor(45, 55, 65))
        set_para_shading(p, CODE_FILL)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def set_para_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_image(doc, src, alt=""):
    path = (SOURCE.parent / src).resolve()
    if not path.exists():
        p = doc.add_paragraph()
        p.add_run(f"[이미지 누락: {src}]").italic = True
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(CONTENT_WIDTH_IN))
    if alt:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(2)
        cap.paragraph_format.space_after = Pt(8)
        r = cap.add_run(alt)
        set_run_font(r, size=9, color=MUTED, italic=True)


def extract_headings(lines):
    headings = []
    for line in lines:
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = line[level:].strip()
            if level <= 2 and text:
                headings.append((level, text.replace(" 초안", "")))
    return headings[1:]


def build():
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    setup_document(doc)

    title = lines[0].lstrip("#").strip()
    add_title_page(doc, title)
    add_toc(doc, extract_headings(lines))

    i = 1
    in_code = False
    code_lines = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        # Markdown tables
        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            headers = split_table_row(stripped)
            rows = []
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_table_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"
            table.autofit = False
            set_table_width(table)
            set_cell_margins(table)
            for row in table.rows:
                prevent_row_split(row)
            for c, header in enumerate(headers):
                cell = table.rows[0].cells[c]
                set_cell_shading(cell, TABLE_HEADER_FILL)
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                add_inline_runs(p, header, bold_default=True)
            for r_idx, row in enumerate(rows, start=1):
                for c_idx, value in enumerate(row[: len(headers)]):
                    cell = table.rows[r_idx].cells[c_idx]
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_after = Pt(0)
                    add_inline_runs(p, value)
            doc.add_paragraph()
            continue

        # HTML image
        html_img = re.search(r'<img\s+[^>]*src="([^"]+)"[^>]*?(?:alt="([^"]*)")?[^>]*>', stripped)
        if html_img:
            add_image(doc, html_img.group(1), html_img.group(2) or "")
            i += 1
            continue

        # Markdown image
        md_img = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if md_img:
            add_image(doc, md_img.group(2), md_img.group(1))
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            h_text = stripped[level:].strip().replace(" 초안", "")
            if level == 2 and re.match(r"^\d+\. ", h_text):
                doc.add_section(WD_SECTION.NEW_PAGE)
            p = doc.add_paragraph()
            p.style = f"Heading {min(level, 4)}"
            add_inline_runs(p, h_text)
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            add_inline_runs(p, stripped[2:])
            i += 1
            continue

        if re.match(r"^\d+\. ", stripped):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(4)
            add_inline_runs(p, re.sub(r"^\d+\. ", "", stripped))
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline_runs(p, stripped)
        i += 1

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
